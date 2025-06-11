from typing import List, Dict

class ReportGenerator:
    """Collects calculation results and outputs them to PDF or Word."""

    def __init__(self) -> None:
        # structured storage of results as a list of dictionaries
        self.results: List[Dict[str, float]] = []

    def add_result(self, name: str, value: float) -> None:
        """Add a calculation result to the internal list."""
        self.results.append({"name": name, "value": value})

    def to_pdf(self, filename: str) -> None:
        """Write results to a simple PDF table."""
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(0, 10, "Results", ln=True)

        for item in self.results:
            pdf.cell(80, 10, item["name"], border=1)
            pdf.cell(40, 10, str(item["value"]), border=1, ln=True)

        pdf.output(filename)

    def to_word(self, filename: str) -> None:
        """Write results to a simple Word table."""
        from docx import Document

        doc = Document()
        doc.add_heading("Results", level=1)
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Name"
        hdr[1].text = "Value"
        for item in self.results:
            row = table.add_row().cells
            row[0].text = item["name"]
            row[1].text = str(item["value"])
        doc.save(filename)

if __name__ == "__main__":
    rg = ReportGenerator()
    # Example calculations
    rg.add_result("A", 1.23)
    rg.add_result("B", 4.56)

    rg.to_pdf("results.pdf")
    rg.to_word("results.docx")
